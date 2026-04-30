from __future__ import annotations

import io
import json
from pathlib import Path

import pytest
from docx import Document

from src.calculations import calculate
from src.document_export import build_docx
from src.eligibility import Check, EligibilityResult, Verdict
from src.form_schema import GrantApplicationInput
from src.generator import GeneratedDraft

FIXTURES = Path(__file__).parent / "fixtures"


def _load_application(name: str) -> GrantApplicationInput:
    return GrantApplicationInput(**json.loads((FIXTURES / name).read_text()))


def _stub_draft() -> GeneratedDraft:
    return GeneratedDraft(
        project_summary=(
            "Novel Hessian-Diagonal Optimiser for Sparse Transformers\n\n"
            "Luminary ML GmbH is developing a curvature-aware optimisation algorithm.\n\n"
            "The team does not know at the project's outset whether the approximation will remain stable."
        ),
        technical_uncertainty=(
            "Existing first-order optimisers exhibit divergence above 85% gradient sparsity. "
            "The specific failure mode arises from unconstrained step sizes in sparse-gradient regions.\n\n"
            "The central technical uncertainty is whether a Hessian-diagonal approximation can remain "
            "numerically stable under these conditions without prohibitive compute overhead."
        ),
        qualifying_activities=(
            "1. Systematic design of controlled training runs to characterise gradient sparsity thresholds. (Experimental development)\n"
            "2. Derivation and implementation of a tractable Hessian-diagonal approximation. (Applied research)\n"
            "3. Ablation studies isolating the contribution of each update rule component. (Experimental development)"
        ),
        consultant_notes=(
            "- [Verify] Annual revenue was not provided — credit rate defaulted to 25% (non-SME).\n"
            "- [Weakness] The project summary does not name specific baseline models used in ablations.\n"
            "- [Strengthen] Ask the client for documented failed training runs to support the uncertainty claim."
        ),
    )


def _all_text(doc: Document) -> str:
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


# ── Test 1: Valid docx bytes (ZIP magic) ─────────────────────────────────────

def test_returns_valid_docx_bytes():
    application = _load_application("clear_rd_input.json")
    calc = calculate(application)
    result = build_docx(application, calc, _stub_draft())
    assert isinstance(result, bytes)
    assert len(result) > 0
    assert result[:4] == b"PK\x03\x04", "Expected ZIP/docx magic bytes"


# ── Test 2: Company name and credit figure appear in document ─────────────────

def test_contains_company_name_and_credit():
    application = _load_application("clear_rd_input.json")
    calc = calculate(application)
    result = build_docx(application, calc, _stub_draft())
    doc = Document(io.BytesIO(result))
    full_text = _all_text(doc)
    assert "Luminary ML GmbH" in full_text
    assert f"€{calc.indicative_credit:,.0f}" in full_text


# ── Test 3: Warnings section absent when no flags fire ────────────────────────

def test_no_warnings_section_when_clean():
    application = _load_application("clear_rd_input.json")
    calc = calculate(application)
    # clear_rd_input has revenue, no capex, not capped — no warnings expected
    assert not calc.revenue_unknown
    assert not calc.capex_excluded
    assert not calc.is_capped
    result = build_docx(application, calc, _stub_draft())
    doc = Document(io.BytesIO(result))
    full_text = _all_text(doc)
    assert "Warnings" not in full_text


# ── Test 4: Warnings section present when revenue_unknown ────────────────────

def test_warnings_section_present_when_revenue_unknown():
    data = json.loads((FIXTURES / "clear_rd_input.json").read_text())
    data["annual_revenue_eur"] = None
    application = GrantApplicationInput(**data)
    calc = calculate(application)
    assert calc.revenue_unknown
    result = build_docx(application, calc, _stub_draft())
    doc = Document(io.BytesIO(result))
    full_text = _all_text(doc)
    assert "Warnings" in full_text
    assert "non-SME" in full_text


# ── Test 5: Qualifying activities render as List Number style ────────────────

def test_numbered_list_style_for_qualifying_activities():
    application = _load_application("clear_rd_input.json")
    calc = calculate(application)
    result = build_docx(application, calc, _stub_draft())
    doc = Document(io.BytesIO(result))
    list_number_paras = [p for p in doc.paragraphs if p.style.name == "List Number"]
    assert len(list_number_paras) >= 1, "Expected at least one List Number paragraph"
    combined = " ".join(p.text for p in list_number_paras)
    assert "Hessian-diagonal" in combined


# ── Test 6: Consultant notes render as List Bullet style ─────────────────────

def test_bullet_style_for_consultant_notes():
    application = _load_application("clear_rd_input.json")
    calc = calculate(application)
    result = build_docx(application, calc, _stub_draft())
    doc = Document(io.BytesIO(result))
    bullet_paras = [p for p in doc.paragraphs if p.style.name == "List Bullet"]
    assert len(bullet_paras) >= 1, "Expected at least one List Bullet paragraph"
    combined = " ".join(p.text for p in bullet_paras)
    assert "[Verify]" in combined or "[Weakness]" in combined or "[Strengthen]" in combined


# ── Test 7: Eligibility section absent when eligibility=None ─────────────────

def test_no_eligibility_section_when_not_passed():
    application = _load_application("clear_rd_input.json")
    calc = calculate(application)
    result = build_docx(application, calc, _stub_draft(), eligibility=None)
    doc = Document(io.BytesIO(result))
    full_text = _all_text(doc)
    assert "Eligibility assessment" not in full_text


# ── Test 8: Eligibility section present when passed ──────────────────────────

def test_eligibility_section_present_when_passed():
    application = _load_application("clear_rd_input.json")
    calc = calculate(application)
    stub_eligibility = EligibilityResult(
        verdict=Verdict.LIKELY_INELIGIBLE,
        checks=[
            Check("German taxable presence", "fail", "No German entity found."),
            Check("R&D classification", "pass", "Clear R&D."),
        ],
        rd_score=5,
        rd_reasoning="Strong research indicators.",
    )
    result = build_docx(application, calc, _stub_draft(), eligibility=stub_eligibility)
    doc = Document(io.BytesIO(result))
    full_text = _all_text(doc)
    assert "Eligibility assessment" in full_text
    assert "Likely Ineligible" in full_text
    assert "[FAIL]" in full_text
    assert "5/5" in full_text
