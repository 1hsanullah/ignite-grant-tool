from __future__ import annotations

import io
import re
from datetime import date
from typing import Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from src.calculations import CostCalculation
from src.eligibility import EligibilityResult
from src.form_schema import GrantApplicationInput
from src.generator import GeneratedDraft

_STATUS_LABEL = {"pass": "[PASS]", "warn": "[WARN]", "fail": "[FAIL]"}

_DISCLAIMER = (
    "Ignite Grant Application Tool — prototype. "
    "Generated content is a drafting aid for qualified consultants; "
    "it is not legal or tax advice."
)


def _add_bold_run(para, text: str, bold: bool) -> None:
    run = para.add_run(text)
    run.bold = bold


def _render_line(para, line: str) -> None:
    """Write a line into para, honoring **bold** spans."""
    parts = re.split(r"\*\*(.+?)\*\*", line)
    for i, part in enumerate(parts):
        if part:
            _add_bold_run(para, part, bold=(i % 2 == 1))


def _render_markdown(doc: Document, text: str) -> None:
    """Render a narrow markdown subset (bullets, numbered lists, paragraphs, **bold**)."""
    blocks = re.split(r"\n{2,}", text.strip())
    numbered_re = re.compile(r"^\d+\.\s+")
    bullet_re = re.compile(r"^[-*]\s+")

    for block in blocks:
        lines = [ln for ln in block.splitlines() if ln.strip()]
        if not lines:
            continue

        first = lines[0]
        if bullet_re.match(first):
            for line in lines:
                content = bullet_re.sub("", line, count=1)
                para = doc.add_paragraph(style="List Bullet")
                _render_line(para, content)
        elif numbered_re.match(first):
            for line in lines:
                content = numbered_re.sub("", line, count=1)
                para = doc.add_paragraph(style="List Number")
                _render_line(para, content)
        else:
            para = doc.add_paragraph()
            _render_line(para, block.replace("\n", " "))


def _add_two_col_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    for i, (label, value) in enumerate(rows):
        table.cell(i, 0).text = label
        table.cell(i, 1).text = value


def _add_cost_breakdown_table(doc: Document, calc: CostCalculation) -> None:
    capex_label = (
        "Capex (excluded — pre-2024)" if calc.capex_excluded else "Capex (100% eligible from 2024)"
    )
    rows = [
        ("Personnel (100% eligible)", f"€{calc.eligible_personnel:,.0f}"),
        ("Contractors (60% eligible)", f"€{calc.eligible_contractor:,.0f}"),
        (capex_label, f"€{calc.eligible_capex:,.0f}"),
        ("Total eligible (before cap)", f"€{calc.total_eligible_before_cap:,.0f}"),
        ("Total eligible (after cap)", f"€{calc.total_eligible:,.0f}"),
    ]
    table = doc.add_table(rows=len(rows) + 1, cols=2)
    table.style = "Table Grid"

    hdr = table.rows[0].cells
    hdr[0].text = "Cost category"
    hdr[1].text = "Amount (€)"
    for cell in hdr:
        for run in cell.paragraphs[0].runs:
            run.bold = True

    for i, (label, value) in enumerate(rows, start=1):
        row = table.rows[i].cells
        is_total = label.startswith("Total")
        row[0].text = label
        row[1].text = value
        if is_total:
            for cell in row:
                for run in cell.paragraphs[0].runs:
                    run.bold = True


def build_docx(
    application: GrantApplicationInput,
    calc: CostCalculation,
    draft: GeneratedDraft,
    eligibility: Optional[EligibilityResult] = None,
) -> bytes:
    doc = Document()

    # ── Title ────────────────────────────────────────────────────────────────
    title_para = doc.add_paragraph(style="Title")
    title_para.add_run(f"FZlG Grant Application Draft — {application.company_name}")

    subtitle = doc.add_paragraph()
    subtitle.add_run(
        f"Claim year {application.claim_year} · Generated {date.today().isoformat()}"
    ).italic = True

    # ── Eligibility assessment (when available) ───────────────────────────────
    if eligibility is not None:
        doc.add_paragraph()
        doc.add_heading("Eligibility assessment", level=1)
        verdict_para = doc.add_paragraph()
        verdict_run = verdict_para.add_run(f"Verdict: {eligibility.verdict.value}")
        verdict_run.bold = True
        for check in eligibility.checks:
            label = _STATUS_LABEL[check.status]
            para = doc.add_paragraph(style="List Bullet")
            para.add_run(f"{label} {check.name}: ").bold = True
            para.add_run(check.reasoning)
        rd_para = doc.add_paragraph()
        rd_para.add_run(f"R&D classifier score: {eligibility.rd_score}/5  ").bold = True
        rd_para.add_run(eligibility.rd_reasoning)

    # ── Indicative cost calculation ───────────────────────────────────────────
    doc.add_heading("Indicative cost calculation", level=1)
    sme_label = "SME (35%)" if calc.is_sme else "Non-SME (25%)"
    if calc.revenue_unknown:
        sme_label += " — unconfirmed"
    _add_two_col_table(doc, [
        ("SME status", sme_label),
        ("Credit rate", f"{calc.credit_rate:.0%}"),
        ("Total eligible costs", f"€{calc.total_eligible:,.0f}"),
        ("Indicative tax credit", f"€{calc.indicative_credit:,.0f}"),
    ])

    # ── Cost breakdown ────────────────────────────────────────────────────────
    doc.add_paragraph()
    doc.add_heading("Cost breakdown", level=1)
    _add_cost_breakdown_table(doc, calc)

    # ── Warnings (only when flags fire) ──────────────────────────────────────
    warnings: list[str] = []
    if calc.revenue_unknown:
        warnings.append(
            "Annual revenue not provided — defaulting to non-SME rate (25%). "
            "If the company qualifies as an SME (<€35M revenue), the credit rate is 35%. "
            "Please verify before submission."
        )
    if calc.capex_excluded:
        warnings.append(
            f"Capital expenditure of €{application.capex_cost_eur:,.0f} was entered "
            f"but claim year {application.claim_year} is before 2024 — capex is excluded from eligible costs."
        )
    if calc.is_capped:
        warnings.append(
            f"Total eligible costs (€{calc.total_eligible_before_cap:,.0f}) exceed the annual cap "
            f"of €{3_500_000:,.0f}. Credit calculated on the capped amount."
        )
    if warnings:
        doc.add_paragraph()
        doc.add_heading("Warnings", level=1)
        for w in warnings:
            doc.add_paragraph(w, style="List Bullet")

    # ── Project summary ───────────────────────────────────────────────────────
    doc.add_paragraph()
    doc.add_heading("Project summary", level=1)
    lines = draft.project_summary.strip().splitlines()
    if lines:
        doc.add_heading(lines[0], level=2)
        remainder = "\n".join(lines[1:]).strip()
        if remainder:
            _render_markdown(doc, remainder)

    # ── Statement of technical uncertainty ────────────────────────────────────
    doc.add_paragraph()
    doc.add_heading("Statement of technical uncertainty", level=1)
    _render_markdown(doc, draft.technical_uncertainty)

    # ── Qualifying R&D activities ─────────────────────────────────────────────
    doc.add_paragraph()
    doc.add_heading("Qualifying R&D activities", level=1)
    _render_markdown(doc, draft.qualifying_activities)

    # ── Notes for the consultant ──────────────────────────────────────────────
    doc.add_paragraph()
    doc.add_heading("Notes for the consultant", level=1)
    doc.add_paragraph(
        "Deterministic flags (from cost inputs) appear first, followed by LLM-identified "
        "weaknesses and items to verify. Review all before submission."
    ).italic = True
    _render_markdown(doc, draft.consultant_notes)

    # ── Disclaimer ────────────────────────────────────────────────────────────
    doc.add_paragraph()
    disc = doc.add_paragraph()
    run = disc.add_run(_DISCLAIMER)
    run.italic = True
    run.font.size = Pt(9)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
